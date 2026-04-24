#!/usr/bin/env python3
"""
Minimal Markdown -> DOCX generator for this repo's docs.

It supports:
- # / ## / ### headings
- bullet lists (- / * )
- numbered lists (1. ...)
- normal paragraphs

This is intentionally lightweight (no pandoc dependency).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: scripts/generate_docx_from_md.py <input.md> <output.docx>")
        return 2

    md_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()

    try:
        from docx import Document  # type: ignore
    except Exception as e:  # noqa: BLE001
        print("Missing dependency: python-docx. Install with: python3 -m pip install --user python-docx")
        print(f"Import error: {e}")
        return 1

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    doc.core_properties.title = md_path.stem

    bullet_re = re.compile(r"^\s*[-*]\s+(.+)\s*$")
    number_re = re.compile(r"^\s*(\d+)\.\s+(.+)\s*$")
    heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*$")

    def add_para(s: str) -> None:
        doc.add_paragraph(s)

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            # Paragraph break: add nothing (Word already separates paragraphs)
            continue

        m = heading_re.match(line)
        if m:
            level = len(m.group(1))
            title = m.group(2)
            # Map markdown heading levels to Word heading styles.
            # Word supports Heading 1..9; clamp to 1..4 for readability.
            style_level = min(max(level, 1), 4)
            doc.add_heading(title, level=style_level)
            continue

        m = bullet_re.match(line)
        if m:
            doc.add_paragraph(m.group(1), style="List Bullet")
            continue

        m = number_re.match(line)
        if m:
            doc.add_paragraph(m.group(2), style="List Number")
            continue

        # Strip simple markdown emphasis markers for docx readability.
        cleaned = line
        cleaned = cleaned.replace("**", "")
        cleaned = cleaned.replace("`", "")
        cleaned = cleaned.replace("> ", "")
        add_para(cleaned)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

